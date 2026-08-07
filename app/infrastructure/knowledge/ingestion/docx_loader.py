"""DOCX loader adapter (python-docx)."""

from __future__ import annotations

import io

from app.intelligence.knowledge.ingestion.ports import DocumentLoaderPort
from app.intelligence.knowledge.models import DocumentFormat, LoadedDocument, RawDocument
from app.shared.exceptions import ValidationFailedError


class DocxDocumentLoader(DocumentLoaderPort):
    @property
    def supported_formats(self) -> set[DocumentFormat]:
        return {DocumentFormat.DOCX}

    async def load(self, document: RawDocument) -> LoadedDocument:
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover
            raise ValidationFailedError("python-docx is required for DOCX ingestion") from exc

        payload = (
            document.content
            if isinstance(document.content, bytes)
            else document.content.encode("utf-8")
        )
        doc = Document(io.BytesIO(payload))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        text = "\n\n".join(paragraphs)
        author = None
        if doc.core_properties and doc.core_properties.author:
            author = doc.core_properties.author
        return LoadedDocument(
            text=text,
            format=DocumentFormat.DOCX,
            source=document.source,
            raw_metadata={"author": author},
        )
